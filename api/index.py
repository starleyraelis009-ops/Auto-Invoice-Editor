from flask import Flask, render_template, request, send_file
from openpyxl import load_workbook
from docx import Document

import os
import re
import tempfile
import zipfile


app = Flask(
    __name__,
    template_folder="../templates"
)


# =========================================================
# NORMALISASI NAMA KOLOM
# =========================================================

def normalize_name(text):

    text = str(text).strip().lower()

    text = text.replace(" ", "_")

    return text


# =========================================================
# GANTI PLACEHOLDER
# =========================================================

def replace_text(text, data):

    pattern = r"\{\{\s*([a-zA-Z0-9_\- ]+)\s*\}\}"


    def replacement(match):

        key = normalize_name(
            match.group(1)
        )

        value = data.get(
            key,
            ""
        )

        if value is None:

            return ""

        return str(value)


    return re.sub(
        pattern,
        replacement,
        text
    )


# =========================================================
# PROSES PARAGRAPH
# =========================================================

def process_paragraph(
    paragraph,
    data
):

    if not paragraph.text:

        return


    new_text = replace_text(
        paragraph.text,
        data
    )


    if new_text != paragraph.text:

        paragraph.text = new_text


# =========================================================
# PROSES WORD
# =========================================================

def process_document(
    template_path,
    output_path,
    data
):

    document = Document(
        template_path
    )


    # Paragraph

    for paragraph in document.paragraphs:

        process_paragraph(
            paragraph,
            data
        )


    # Table

    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    process_paragraph(
                        paragraph,
                        data
                    )


    # Header dan footer

    for section in document.sections:

        for paragraph in section.header.paragraphs:

            process_paragraph(
                paragraph,
                data
            )


        for paragraph in section.footer.paragraphs:

            process_paragraph(
                paragraph,
                data
            )


    document.save(
        output_path
    )


# =========================================================
# HALAMAN UTAMA
# =========================================================

@app.route("/")
def home():

    return render_template(
        "index.html"
    )


# =========================================================
# GENERATE
# =========================================================

@app.route(
    "/generate",
    methods=["POST"]
)
def generate():

    excel = request.files.get(
        "excel"
    )

    template = request.files.get(
        "template"
    )


    if not excel or not template:

        return (
            "Excel dan template Word wajib dipilih.",
            400
        )


    if not excel.filename.lower().endswith(
        ".xlsx"
    ):

        return (
            "File Excel harus .xlsx.",
            400
        )


    if not template.filename.lower().endswith(
        ".docx"
    ):

        return (
            "Template Word harus .docx.",
            400
        )


    # =====================================================
    # TEMPORARY DIRECTORY
    # =====================================================

    workdir = tempfile.mkdtemp()


    try:

        excel_path = os.path.join(
            workdir,
            "data.xlsx"
        )


        template_path = os.path.join(
            workdir,
            "template.docx"
        )


        excel.save(
            excel_path
        )


        template.save(
            template_path
        )


        # =================================================
        # BACA EXCEL
        # =================================================

        workbook = load_workbook(
            excel_path,
            data_only=True
        )


        worksheet = workbook.active


        headers = []


        for cell in worksheet[1]:

            if cell.value is not None:

                headers.append(
                    normalize_name(
                        cell.value
                    )
                )


        if not headers:

            return (
                "Excel tidak memiliki header.",
                400
            )


        # =================================================
        # BACA DATA
        # =================================================

        rows = []


        for row in worksheet.iter_rows(
            min_row=2,
            values_only=True
        ):


            if not any(
                value is not None
                for value in row
            ):

                continue


            data = {}


            for index, header in enumerate(
                headers
            ):

                if index < len(row):

                    data[header] = row[index]

                else:

                    data[header] = ""


            rows.append(
                data
            )


        if not rows:

            return (
                "Excel tidak memiliki data.",
                400
            )


        # =================================================
        # HASIL
        # =================================================

        output_dir = os.path.join(
            workdir,
            "hasil"
        )


        os.makedirs(
            output_dir,
            exist_ok=True
        )


        generated_files = []


        # =================================================
        # GENERATE SETIAP BARIS
        # =================================================

        for number, data in enumerate(
            rows,
            start=1
        ):

            name = data.get(
                "nama",
                f"data_{number}"
            )


            safe_name = re.sub(
                r"[^a-zA-Z0-9_\-]",
                "_",
                str(name)
            )


            filename = (
                f"{number}_{safe_name}.docx"
            )


            output_path = os.path.join(
                output_dir,
                filename
            )


            process_document(
                template_path,
                output_path,
                data
            )


            generated_files.append(
                output_path
            )


        # =================================================
        # SATU FILE
        # =================================================

        if len(generated_files) == 1:

            return send_file(
                generated_files[0],
                as_attachment=True,
                download_name=os.path.basename(
                    generated_files[0]
                )
            )


        # =================================================
        # BANYAK FILE → ZIP
        # =================================================

        zip_path = os.path.join(
            workdir,
            "hasil_dokumen.zip"
        )


        with zipfile.ZipFile(
            zip_path,
            "w",
            zipfile.ZIP_DEFLATED
        ) as zip_file:

            for file_path in generated_files:

                zip_file.write(
                    file_path,
                    arcname=os.path.basename(
                        file_path
                    )
                )


        return send_file(
            zip_path,
            as_attachment=True,
            download_name="hasil_dokumen.zip"
        )


    except Exception as error:

        print(
            "ERROR:",
            error
        )


        return (
            f"Terjadi kesalahan: {error}",
            500
        )


# =========================================================
# VERCEL
# =========================================================

if __name__ == "__main__":

    app.run(
        host="0.0.0.0",
        port=5000,
        debug=True
    )